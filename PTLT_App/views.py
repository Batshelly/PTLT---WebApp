from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.views.decorators.csrf import csrf_protect
import json
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.template.loader import render_to_string
from django.db.models import Q
from django.core.paginator import Paginator
from django.views.decorators.http import require_http_methods   
import json
from django.contrib.auth.models import User
from django.utils.crypto import get_random_string
from django.db import transaction
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.contrib.auth import logout
from django.utils.safestring import mark_safe
from django.core.serializers.json import DjangoJSONEncoder
from datetime import time, timedelta
from django.views.decorators.http import require_POST
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_str
from django.contrib.auth.hashers import check_password
from django.contrib.auth.tokens import default_token_generator
from django.contrib.sites.shortcuts import get_current_site
from django.db import IntegrityError
from django.contrib.auth import login as auth_login
import datetime
import csv
import io
import traceback
from django.core.mail import send_mail
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.utils import timezone
from datetime import datetime, date
from functools import wraps
import PyPDF2
import re
from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.authentication import TokenAuthentication
from rest_framework.permissions import IsAuthenticated
from django.contrib.auth import authenticate
from rest_framework.authtoken.models import Token
from rest_framework import serializers
from django.contrib.auth.hashers import make_password
# Authentication endpoint for mobile
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from .serializers import MobileClassScheduleSerializer
from PTLT_App.utils import archive_semester_data
from .models import Account
from .models import CourseSection
from .models import ClassSchedule
from .models import AttendanceRecord
from .models import Semester
from .models import AccountUploadNotification

from collections import defaultdict
from django.utils.dateparse import parse_date


from .serializers import (
    AccountSerializer, ClassScheduleSerializer, AttendanceRecordSerializer,
    MobileAccountSerializer, MobileAttendanceSerializer, CourseSectionSerializer  
)


# for docx file
from docxtpl import DocxTemplate
from io import BytesIO
import os
from django.conf import settings
from django.core.mail import EmailMessage
import random
from datetime import datetime, timedelta, date
from .models import Account, CourseSection, ClassSchedule
import re
import logging
import tempfile
import subprocess
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfMerger





def admin_required(view_func):
    """Decorator for admin-only views - handles both AJAX and regular requests"""
    @wraps(view_func)
    @login_required
    def wrapper(request, *args, **kwargs):
        try:
            account = Account.objects.get(email=request.user.email, role="Admin")
        except Account.DoesNotExist:
            # Check if it's an AJAX request
            if request.headers.get('Content-Type') == 'application/json' or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'status': 'error',
                    'message': 'Access denied: Admin privileges required.'
                }, status=403)
            else:
                messages.error(request, "Access denied: Admin privileges required.")
                return redirect('login')
        
        # ⚠️ CRITICAL: Must return the view function's response
        return view_func(request, *args, **kwargs)
    return wrapper


def instructor_or_admin_required(view_func):
    @wraps(view_func)
    @login_required
    def wrapper(request, *args, **kwargs):
        try:
            account = Account.objects.get(email=request.user.email, role__in=['Instructor', 'Admin'])
        except Account.DoesNotExist:
            messages.error(request, "Access denied: Instructor or Admin role required.")
            return redirect('login')
        return view_func(request, *args, **kwargs)
    return wrapper

def instructor_required(view_func):
    @wraps(view_func)
    @login_required
    def wrapper(request, *args, **kwargs):
        try:
            account = Account.objects.get(email=request.user.email, role='Instructor')
        except Account.DoesNotExist:
            messages.error(request, "Access denied: Instructor role required.")
            return redirect('login')
        return view_func(request, *args, **kwargs)
    return wrapper

#para to sa schedule ni instructor
PERIODS = [
    (time(9, 30), time(10, 20), "I"),
    (time(10, 20), time(11, 10), "II"),
    (time(11, 10), time(12, 0), "III"),
    (time(12, 0), time(12, 40), "Break"),  # Lunch
    (time(12, 40), time(13, 30), "IV"),
    (time(13, 30), time(14, 20), "V"),
    (time(14, 20), time(15, 10), "VI"),
    (time(15, 10), time(16, 0), "VII"),
]

@transaction.atomic 
def login_view(request):
    
    # 🔥 NEW: Clear all messages when session expires
    session_expired = request.GET.get('session_expired', False)
    if session_expired:
        # Clear all existing messages
        from django.contrib.messages import get_messages
        storage = get_messages(request)
        storage.used = True  # Mark all messages as used to clear them
        
        # Add the session expired message
        messages.warning(request, 'Your session has expired. Please log in again.')
        
    # Check if any semester ended and not archived yet
    today = timezone.now().date()
    semester_to_archive = Semester.objects.filter(end_date__lt=today, is_archived=False).first()

    if semester_to_archive:
        print("Will archive now.")
        # Call your archive function
        archived_count = archive_semester_data(semester_to_archive)
        print(f"Archived {archived_count} records.")
    else:
        print("No archive needed at this time.")


    # Check if any accounts exist - keep your existing default account creation
    if not Account.objects.exists():
        default_accounts = [
            {
                'user_id': '000000',
                'email': 'tupcptlt@gmail.com',
                'first_name': 'Super Admin',
                'last_name': 'Account 0',
                'role': 'Admin',
                'password': None,  # No password needed
                'sex': 'Other',
                'status': 'Active'
            },
            {
                'user_id': '000001',
                'email': 'shelwinjay.buenaventura@gsfe.tupcavite.edu.ph',
                'first_name': 'Dummy',
                'last_name': 'Account 1',
                'role': 'Instructor',
                'password': None,
                'sex': 'Other',
                'status': 'Active'
            },
            {
                'user_id': '000002',
                'email': 'marktrieste.milan@gsfe.tupcavite.edu.ph',
                'first_name': 'Dummy',
                'last_name': 'Account 2',
                'role': 'Instructor',
                'password': None,
                'sex': 'Other',
                'status': 'Active'
            },
            {
                'user_id': '000003',
                'email': 'markjoshua.salinas@gsfe.tupcavite.edu.ph',
                'first_name': 'Dummy',
                'last_name': 'Account 3',
                'role': 'Instructor',
                'password': None,
                'sex': 'Other',
                'status': 'Active'
            },
            {
                'user_id': '000004',
                'email': 'janxander.yangco@gsfe.tupcavite.edu.ph',
                'first_name': 'Dummy',
                'last_name': 'Account 4',
                'role': 'Instructor',
                'password': None,
                'sex': 'Other',
                'status': 'Active'
            }
        ]

        for acc in default_accounts:
            # Create Django User for session management
            user = User.objects.create_user(
                username=acc['user_id'],
                email=acc['email'],
                password=get_random_string(12),  # Random password (won't be used)
                first_name=acc['first_name'],
                last_name=acc['last_name']
            )

            # Create Account entry
            Account.objects.create(
                user_id=acc['user_id'],
                email=acc['email'],
                first_name=acc['first_name'],
                last_name=acc['last_name'],
                role=acc['role'],
                password=None,
                sex=acc['sex'],
                status=acc['status'],
                course_section=None
            )
    
    if request.method == 'POST':
        email = request.POST.get('email')

        try:
            # Check if account exists and is Instructor or Admin
            account = Account.objects.get(email=email, role__in=['Instructor', 'Admin'])
            
            # Check if account is active
            if account.status != 'Active':
                messages.error(request, 'Your account is not active. Please contact administrator.')
                return redirect('login')
            
            # Generate 6-digit OTP
            otp = random.randint(100000, 999999)
            
            # Store OTP in session with timestamp
            request.session['login_otp'] = otp
            request.session['login_email'] = email
            request.session['otp_timestamp'] = datetime.now().isoformat()
            
            # Design the email
            email_subject = "PTLT - Login OTP Verification"
            email_body = f"""
            <html>
                <body style="font-family: Arial, sans-serif;">
                    <h2 style="color: #661e1e;">PTLT Login Verification</h2>
                    <p>Hello {account.first_name},</p>
                    <p>You have requested to log in to the PTLT system. Please use the OTP below:</p>
                    <div style="background-color: #f5f5f5; padding: 20px; text-align: center; margin: 20px 0;">
                        <h1 style="color: #661e1e; font-size: 36px; margin: 0; letter-spacing: 5px;">{otp}</h1>
                    </div>
                    <p style="color: #666;">This OTP is valid for <strong>5 minutes</strong>. Please do not share it with anyone.</p>
                    <hr style="border: none; border-top: 1px solid #ddd; margin: 20px 0;">
                    <p style="font-size: 12px; color: #999;">
                        If you did not attempt to log in, please ignore this email or contact the administrator.
                    </p>
                </body>
            </html>
            """

            # Send OTP email
            try:
                email_message = EmailMessage(
                    email_subject,
                    email_body,
                    settings.EMAIL_HOST_USER,
                    [email],
                )
                email_message.content_subtype = 'html'
                email_message.send()
                
                messages.success(request, f'OTP has been sent to {email}. Please check your inbox.')
                return redirect('verify_login_otp')
                
            except Exception as e:
                messages.error(request, 'Failed to send OTP. Please try again later.')
                return redirect('login')

        except Account.DoesNotExist:
            messages.error(request, 'No account found with that email, or you do not have access to this system.')
            return redirect('login')
        
    return render(request, 'login.html')

def verify_login_otp(request):
    """
    Verify OTP and authenticate user.
    Handles dual Account/User table architecture safely.
    """
    # Check if user came from login page with OTP
    if 'login_otp' not in request.session or 'login_email' not in request.session:
        messages.error(request, "Session expired. Please request a new OTP.")
        return redirect('login')
    
    if request.method == 'POST':
        user_otp = request.POST.get('otp')
        stored_otp = request.session.get('login_otp')
        email = request.session.get('login_email')
        otp_timestamp = request.session.get('otp_timestamp')
        
        # Check OTP expiration (5 minutes)
        try:
            otp_time = datetime.fromisoformat(otp_timestamp)
            current_time = datetime.now()
            time_diff = (current_time - otp_time).total_seconds() / 60
            
            if time_diff > 5:
                # OTP expired - clear session
                del request.session['login_otp']
                del request.session['login_email']
                del request.session['otp_timestamp']
                messages.error(request, "OTP has expired. Please request a new one.")
                return redirect('login')
        except (ValueError, TypeError):
            messages.error(request, "Invalid session data. Please try again.")
            return redirect('login')
        
        # Validate OTP
        if str(stored_otp) != user_otp:
            # Invalid OTP - don't clear session, allow retry
            messages.error(request, "Invalid OTP. Please try again.")
            return render(request, 'verify_login_otp.html', {'email': email})
        
        # OTP is valid - proceed with authentication
        try:
            # 1. Get the Account record by email
            account = Account.objects.get(email=email)
            
            # 2. Check account status
            if account.status != 'Active':
                messages.error(request, f"Your account is {account.status}. Please contact admin.")
                return redirect('login')
            
            # 3. Get or create Django User (synchronized to Account)
            try:
                # Try to get existing User by username (which maps to account.user_id)
                user_obj = User.objects.get(username=account.user_id)
                
                # Update User fields to match Account (in case email changed)
                needs_update = False
                if user_obj.email != account.email:
                    user_obj.email = account.email
                    needs_update = True
                if user_obj.first_name != account.first_name:
                    user_obj.first_name = account.first_name
                    needs_update = True
                if user_obj.last_name != account.last_name:
                    user_obj.last_name = account.last_name
                    needs_update = True
                
                if needs_update:
                    user_obj.save()
                    
            except User.DoesNotExist:
                # User doesn't exist - create it
                try:
                    user_obj = User.objects.create_user(
                        username=account.user_id,
                        email=account.email,
                        password=get_random_string(12),  # Random password (not used for login)
                        first_name=account.first_name,
                        last_name=account.last_name
                    )
                except IntegrityError as e:
                    # Handle duplicate username gracefully
                    if 'username' in str(e).lower() or 'duplicate' in str(e).lower():
                        messages.error(request, "Account synchronization error. Please contact admin.")
                    else:
                        messages.error(request, "Authentication failed. Please try again.")
                    return redirect('login')
            
            # 4. Log the user in (backend bypasses password check)
            auth_login(request, user_obj, backend='django.contrib.auth.backends.ModelBackend')
            
            # 5. Set session variables
            request.session['user_id'] = account.user_id
            request.session['role'] = account.role
            
            # 6. Clear OTP data from session
            del request.session['login_otp']
            del request.session['login_email']
            del request.session['otp_timestamp']
            
            messages.success(request, f"Welcome back, {account.first_name}!")
            
            # 7. Redirect based on role
            if account.role == 'Admin':
                return redirect('account_management')
            elif account.role == 'Instructor':
                return redirect('instructor_schedule')
            elif account.role == 'Student':
                return redirect('student_dashboard')  # Adjust to your actual student URL
            else:
                messages.error(request, "Unknown user role.")
                return redirect('login')
                
        except Account.DoesNotExist:
            messages.error(request, "Account not found. Please check your email or contact admin.")
            return redirect('login')
        
        except Exception as e:
            # Catch any unexpected errors
            messages.error(request, "An unexpected error occurred. Please try again.")
            # Log the error for debugging (optional)
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Login error for {email}: {str(e)}")
            return redirect('login')
    
    # GET request - show OTP entry form
    email = request.session.get('login_email')
    return render(request, 'verify_login_otp.html', {'email': email})
# NEW VIEW: Resend OTP
def resend_login_otp(request):
    if 'login_email' not in request.session:
        messages.error(request, "Session expired. Please start login process again.")
        return redirect('login')
    
    email = request.session.get('login_email')
    
    try:
        account = Account.objects.get(email=email, role__in=['Instructor', 'Admin'])
        
        # Generate new OTP
        otp = random.randint(100000, 999999)
        
        # Update session
        request.session['login_otp'] = otp
        request.session['otp_timestamp'] = datetime.now().isoformat()
        
        # Send email
        email_subject = "PTLT - New Login OTP"
        email_body = f"""
        <html>
            <body style="font-family: Arial, sans-serif;">
                <h2 style="color: #661e1e;">PTLT Login Verification</h2>
                <p>Hello {account.first_name},</p>
                <p>You have requested a new OTP. Please use the code below:</p>
                <div style="background-color: #f5f5f5; padding: 20px; text-align: center; margin: 20px 0;">
                    <h1 style="color: #661e1e; font-size: 36px; margin: 0; letter-spacing: 5px;">{otp}</h1>
                </div>
                <p style="color: #666;">This OTP is valid for <strong>5 minutes</strong>.</p>
            </body>
        </html>
        """
        
        email_message = EmailMessage(
            email_subject,
            email_body,
            settings.EMAIL_HOST_USER,
            [email],
        )
        email_message.content_subtype = 'html'
        email_message.send()
        
        messages.success(request, 'A new OTP has been sent to your email.')
        return redirect('verify_login_otp')
        
    except Exception as e:
        messages.error(request, 'Failed to resend OTP. Please try again.')
        return redirect('verify_login_otp')


def force_password_change(request):
    # Check if user came from login with temp password
    if 'temp_user_id' not in request.session:
        messages.error(request, "Unauthorized access.")
        return redirect('login')
    
    if request.method == 'POST':
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')

        if new_password == '000000':
            messages.error(request, "Please try a different password.")
            return render(request, 'force_password_change.html')
        
        if new_password == 'secret':
            messages.error(request, "Uy bat mo alam?")
            return render(request, 'force_password_change.html')
        
        if new_password != confirm_password:
            messages.error(request, "Passwords do not match.")
            return render(request, 'force_password_change.html')
        
        if len(new_password) < 6:  # Add your password requirements
            messages.error(request, "Password must be at least 6 characters long.")
            return render(request, 'force_password_change.html')
        
        try:
            user_id = request.session['temp_user_id']
            email = request.session['temp_email']
            account = Account.objects.get(user_id=user_id, email=email)
            
            # Create or update Django User
            try:
                user_obj = User.objects.get(email=email)
                # Update existing user password
                user_obj.set_password(new_password)
                user_obj.save()
            except User.DoesNotExist:
                # Create new Django User
                user_obj = User.objects.create_user(
                    username=account.user_id,
                    email=account.email,
                    password=new_password,
                    first_name=account.first_name,
                    last_name=account.last_name
                )
            
            # Clear temporary session data
            del request.session['temp_user_id']
            del request.session['temp_email']
            
            # Auto-login the user
            user = authenticate(request, username=user_obj.username, password=new_password)
            if user:
                login(request, user)
                request.session['user_id'] = account.user_id
                request.session['role'] = account.role
                
                messages.success(request, "Password updated successfully!")
                
                # Redirect based on role
                if account.role == 'Admin':
                    return redirect('account_management')
                elif account.role == 'Instructor':
                    return redirect('instructor_schedule')
                else:
                    return redirect('login')
            
        except Account.DoesNotExist:
            messages.error(request, "Account not found.")
            return redirect('login')
    
    return render(request, 'force_password_change.html')

def logout_view(request):
    logout(request)  # Destroys session and logs out user
    return redirect('login') 

@admin_required
def create_instructor(request):
    if request.method == 'POST':
        form_type = request.POST.get('form_type')

        if form_type == 'instructor_form':
            # Instructor form submitted
            user_id = request.POST.get('user_id')
            email = request.POST.get('email')
            first_name = request.POST.get('first_name')
            last_name = request.POST.get('last_name')
            role = request.POST.get('role')
            password = request.POST.get('password')
            sex = request.POST.get('sex')

            if Account.objects.filter(user_id=user_id).exists():
                messages.error(request, 'User ID already exists.')
            elif Account.objects.filter(email=email).exists():
                messages.error(request, 'Email already exists.')
            else:
                Account.objects.create(
                    user_id=user_id,
                    email=email,
                    first_name=first_name,
                    last_name=last_name,
                    role=role,
                    password=password,
                    sex=sex
                )
                messages.success(request, 'Instructor account created successfully!')
                return redirect('create_instructor')

        elif form_type == 'course_section_form':
            # Course & section form submitted
            course_name = request.POST.get('course_name')
            section_name = request.POST.get('section_name')

            try:
                course_section = CourseSection.objects.create(
                    course_name=course_name,
                    section_name=section_name
                )
                messages.success(request, f'Course Section "{course_section.course_section}" created successfully!')
                return redirect('create_instructor')
            except Exception as e:
                messages.error(request, f'Error: {str(e)}')

    return render(request, 'create_instructor.html')

def forgot_password(request):
    if request.method == "POST":
        email = request.POST.get('email')

        try:
            # Fetch the user using the default User model
            user = User.objects.get(email=email)

            # Generate the password reset token
            token = default_token_generator.make_token(user)
            uid = urlsafe_base64_encode(force_bytes(user.pk))
            current_site = get_current_site(request).domain
            # Generate the reset password URL
            reset_link = f"http://{current_site}/reset-password/{uid}/{token}/"

            # HTML email content - using triple quotes without f-string for CSS, then formatting
            email_subject = 'Password Reset Request'
            
            # Create the email body using format() method instead of f-string
            email_body = """
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Password Reset</title>
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        margin: 0;
                        padding: 0;
                        background-color: #f5f5f5;
                    }}
                    .container {{
                        width: 100%;
                        max-width: 600px;
                        margin: 0 auto;
                        background-color: #ffffff;
                        padding: 20px;
                        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
                    }}
                    h1 {{
                        color: #333;
                        text-align: center;
                    }}
                    p {{
                        font-size: 1rem;
                        line-height: 1.5;
                        color: #555;
                    }}
                    .button {{
                        display: inline-block;
                        padding: 12px 24px;
                        background-color: #661e1e;
                        color: white !important;
                        text-decoration: none;
                        border-radius: 4px;
                        font-size: 1rem;
                        margin-top: 20px;
                        text-align: center;
                        transition: background-color 0.3s ease;
                    }}
                    .button:hover {{
                        background-color: #a74545;
                    }}
                    .footer {{
                        margin-top: 30px;
                        padding-top: 20px;
                        border-top: 1px solid #eee;
                        font-size: 0.9rem;
                        color: #777;
                    }}
                    .warning {{
                        background-color: #fff3cd;
                        border: 1px solid #ffeaa7;
                        border-radius: 4px;
                        padding: 15px;
                        margin: 20px 0;
                        color: #856404;
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header-accent"></div>
                    <h1>Password Reset Request</h1>
                    <p>Hello {first_name},</p>
                    <p>We received a request to reset your password for your account. To proceed with resetting your password, please click the button below:</p>
                    <p><a href="{reset_link}" class="button">Reset Password</a></p>
                    
                    <div class="warning">
                        <strong>Security Notice:</strong> This link will expire in 24 hours for your security. If you didn't request this password reset, please ignore this email and your password will remain unchanged.
                    </div>
                    
                    <div class="footer">
                        <p>If the button above doesn't work, you can copy and paste this link into your browser:</p>
                        <p style="word-break: break-all; color: #dc2626 !important; font-weight: 600; text-decoration: none !important;">{reset_link}</p>
                        <p style="margin-top: 20px;">
                            <span style="color: #6b7280;">Best regards,</span><br>
                            <strong style="color: #dc2626;">PTLT TUP-CAVITE</strong>
                        </p>
                    </div>
                </div>
            </body>
            </html>
            """.format(first_name=user.first_name, reset_link=reset_link)

            # Send the HTML email
            send_mail(
                email_subject,
                '',  # Plain text version of the email (empty since we are sending HTML)
                'from@example.com',  # Set your sender email
                [email],
                fail_silently=False,
                html_message=email_body  # HTML version of the email
            )

            # Show success message to the user
            messages.success(request, 'Password reset link has been sent to your email address.')
            return redirect('login')

        except User.DoesNotExist:
            # Handle the case where the user doesn't exist
            messages.error(request, 'No account found with this email address.')
    return render(request, 'forgot_password.html')

def reset_password(request, encoded_email, token):
    try:
        # Decode the user ID from the encoded email
        try:
            uid = urlsafe_base64_decode(encoded_email).decode('utf-8')
            print(f"Decoded user ID: {uid}")
        except Exception as e:
            print(f"Error decoding email: {e}")
            messages.error(request, 'Invalid or expired reset link.')
            return redirect('login')
        
        # Fetch the user based on the decoded ID
        try:
            user = User.objects.get(pk=uid)
        except User.DoesNotExist:
            print(f"User does not exist for ID: {uid}")
            messages.error(request, 'User not found.')
            return redirect('login')
        
        # Check if the token matches the user's reset token
        if default_token_generator.check_token(user, token):
            if request.method == 'POST':
                new_password = request.POST.get('password')
                confirm_password = request.POST.get('confirm_password')
                print(f"new_password: {new_password}, confirm_password: {confirm_password}")

                # Ensure the passwords match
                if new_password == confirm_password:
                    user.set_password(new_password)  # Set the new password
                    user.save()  # Save the user object
                    print("Password reset successfully!")
                    messages.success(request, 'Your password has been reset successfully!')
                    update_session_auth_hash(request, user)  # Keep the user logged in
                    return redirect('login')  # Redirect to login page
                else:
                    messages.error(request, 'Passwords do not match. Please try again.')

            return render(request, 'reset_password.html', {'uid': encoded_email, 'token': token})

        else:
            print(f"Invalid token for user: {uid}")
            messages.error(request, 'Invalid or expired reset link.')
            return redirect('login')

    except Exception as e:
        print(f"Error occurred: {e}")
        messages.error(request, 'An error occurred during password reset. Please try again later.')
        return redirect('login')
    

@instructor_required
def student_attendance_records(request):
    # Get logged-in instructor's Account entry
    try:
        instructor_account = Account.objects.get(email=request.user.email, role='Instructor')
    except Account.DoesNotExist:
        return render(request, "error.html", {"message": "Instructor account not found"})

    # Subjects/Courses taught by this instructor
    schedules = ClassSchedule.objects.filter(professor=instructor_account)

    selected_schedule_id = request.GET.get("schedule")
    selected_date_range = request.GET.get("date_range")
    date_ranges = []
    attendance_table = []

    if selected_schedule_id:
        # Fetch unique attendance dates for selected class schedule
        attendance_dates = AttendanceRecord.objects.filter(
            class_schedule_id=selected_schedule_id
        ).values_list("date", flat=True).distinct().order_by("date")

        attendance_dates = list(attendance_dates)

        # Group into 8-day ranges for the filter dropdown
        for i in range(0, len(attendance_dates), 8):
            start_date = attendance_dates[i]
            end_date = attendance_dates[min(i + 7, len(attendance_dates) - 1)]
            date_ranges.append({
                "value": f"{start_date}_to_{end_date}",
                "label": f"{start_date.strftime('%B %d, %Y')} - {end_date.strftime('%B %d, %Y')}"
            })

        if selected_date_range:
            try:
                start_str, end_str = selected_date_range.split("_to_")
                start_date = parse_date(start_str)
                end_date = parse_date(end_str)
            except (ValueError, TypeError):
                start_date = end_date = None

            if start_date and end_date:
                # Get all attendance records within the date range
                attendance_qs = AttendanceRecord.objects.filter(
                    class_schedule_id=selected_schedule_id,
                    date__range=(start_date, end_date)
                ).select_related('student')

                # Get schedule object once
                schedule_obj = ClassSchedule.objects.get(id=selected_schedule_id)

                # Get students in the same course_section as the schedule
                students_in_schedule = Account.objects.filter(
                    course_section_id=schedule_obj.course_section_id
                ).order_by('last_name', 'first_name')

                # Map attendance data: {student_id: {date: {'status': status, 'time_in': time_in, 'time_out': time_out}}}
                attendance_data = defaultdict(lambda: defaultdict(dict))
                for record in attendance_qs:
                    attendance_data[record.student.id][record.date] = {
                        'status': record.status,
                        'time_in': record.time_in,
                        'time_out': record.time_out
                    }

                # Build date headers (max 8 dates)
                date_headers = [d for d in attendance_dates if start_date <= d <= end_date][:8]
                num_empty_date_column = 8 - len(date_headers)
                #create a pseudo list just for the for loop in html to work
                num_empty_date_columns = []
                for i in range(num_empty_date_column):
                    num_empty_date_columns.append("")


                attendance_table = []
                for student in students_in_schedule:
                    course_section_for_student = student.course_section
                    # Build dates_statuses as a list of dicts containing all attendance info
                    dates_statuses = []
                    for date in date_headers:
                        attendance_info = attendance_data[student.id].get(date, {})
                        dates_statuses.append({
                            'status': attendance_info.get('status', ''),
                            'time_in': attendance_info.get('time_in', ''),
                            'time_out': attendance_info.get('time_out', '')
                        })
                    
                    # Pad with empty dicts to always have 8 items
                    while len(dates_statuses) < 8:
                        dates_statuses.append({
                            'status': '',
                            'time_in': '',
                            'time_out': ''
                        })

                    # DEBUG: Check the length
                    print(f"Student {student.user_id}: dates_statuses length = {len(dates_statuses)}")

                    row = {
                        "student_id": student.user_id, 
                        "name": f"{student.first_name} {student.last_name}",
                        "sex": student.sex,
                        "course": course_section_for_student,
                        "subject": schedule_obj.course_code,
                        "room": schedule_obj.room_assignment,
                        "dates": dates_statuses,  # Always 8 items now
                    }
                    attendance_table.append(row)

                context = {
                    "schedules": schedules,
                    "date_ranges": date_ranges,
                    "selected_schedule_id": selected_schedule_id,
                    "selected_date_range": selected_date_range,
                    "attendance_table": attendance_table,
                    "date_headers": date_headers,
                    "num_empty_date_columns": num_empty_date_columns,
                }
                return render(request, "student_attendance_records.html", context)

        # If no valid date range selected or dates invalid, still render with schedules & date ranges
        context = {
            "schedules": schedules,
            "date_ranges": date_ranges,
            "selected_schedule_id": selected_schedule_id,
            "attendance_table": [],
            "date_headers": [],
            "num_empty_date_columns": ["", "", "", "", "", "", "", "", ],
        }
        return render(request, "student_attendance_records.html", context)

    # If no schedule selected at all, render with just schedules and empty date ranges
    context = {
        "schedules": schedules,
        "date_ranges": date_ranges,
        "selected_schedule_id": selected_schedule_id,
        "attendance_table": [],
        "date_headers": [],
        "num_empty_date_columns": ["", "", "", "", "", "", "", "", ],
    }
    return render(request, "student_attendance_records.html", context)



@require_POST
@instructor_required
def update_class_schedule_instructor(request):
    try:
        data = json.loads(request.body)

        # Validate required fields
        required_fields = ['course_code', 'time_in', 'time_out', 'room_assignment', 'grace_period']
        for field in required_fields:
            if field not in data:
                return JsonResponse({'success': False, 'error': f'Missing field: {field}'}, status=400)

        # Match instructor via user_id
        instructor = Account.objects.get(user_id=request.user.username, role='Instructor')

        # Get the specific schedule (only one per instructor + course_code assumed)
        schedule = ClassSchedule.objects.get(course_code=data['course_code'], professor=instructor)

        # Apply updates
        schedule.time_in = data['time_in']
        schedule.time_out = data['time_out']
        schedule.room_assignment = data['room_assignment']
        schedule.grace_period = int(data['grace_period'])
        schedule.save()

        return JsonResponse({'success': True})

    except ClassSchedule.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Class schedule not found.'}, status=404)

    except Account.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Instructor account not found.'}, status=403)

    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON data.'}, status=400)

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@instructor_required    
def instructor_schedule(request):
    user = request.user
    try:
        instructor = Account.objects.get(user_id=user.username, role='Instructor')
    except Account.DoesNotExist:
        return render(request, 'error.html', {'message': 'Instructor not found.'})

    schedules = ClassSchedule.objects.filter(professor=instructor)

    for schedule in schedules:
        student_acc = Account.objects.filter(course_section_id=schedule.course_section_id)
        student_count = len(student_acc)
        schedule.student_count = int(student_count)
        schedule.save()

    return render(request, 'schedule.html', {
        'class_schedules': schedules,
    })


@admin_required
def account_management(request):
    role_filter = request.GET.get('role', '')
    status_filter = request.GET.get('status', '')
    course_filter = request.GET.get('course', '')
    search_query = request.GET.get('search', '')
    
    accounts = Account.objects.all()
    
    if role_filter:
        accounts = accounts.filter(role__iexact=role_filter)
    if status_filter:
        accounts = accounts.filter(status__iexact=status_filter)
    if course_filter:
        accounts = accounts.filter(course_section__course_section=course_filter)
    if search_query:
        accounts = accounts.filter(
            Q(first_name__icontains=search_query) | Q(last_name__icontains=search_query)
        )
    
    # Order the accounts
    accounts = accounts.order_by('user_id')
    
    # Get all course sections for dropdown
    course_sections = CourseSection.objects.all().order_by('course_section')
    
    # Add pagination - 10 accounts per page
    paginator = Paginator(accounts, 10)
    page_number = request.GET.get('page')
    accounts_page = paginator.get_page(page_number)
    
    update_notifications = AccountUploadNotification.objects.filter(is_read=False, notification_type='update')
    update_count = update_notifications.count()
    recent_updates = update_notifications[:5]
    
    # Mark updates as read if requested
    if request.GET.get('mark_updates_read') == 'true':
        AccountUploadNotification.objects.filter(is_read=False, notification_type='update').update(is_read=True)
        messages.success(request, f'Marked {update_count} notifications as read.')
        return redirect('account_management')
    
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        html = render_to_string('partials/account_table_body.html', {'accounts': accounts_page})
        return JsonResponse({'html': html})
    
    # Build query string for pagination links
    query_params = request.GET.copy()
    if 'page' in query_params:
        query_params.pop('page')
    query_string = query_params.urlencode()
    
    context = {
        'accounts': accounts_page,
        'update_count': update_count,
        'query_string': query_string,
        'course_sections': course_sections,
        'role_filter': role_filter,         
        'status_filter': status_filter,     
        'course_filter': course_filter,     
        'search_query': search_query,       
    }
    return render(request, 'account_management.html', context)

@csrf_exempt
@admin_required
@require_http_methods(["POST"])
def toggle_account_status(request, account_id):
    """
    Toggle account status between Active and Inactive
    """
    try:
        account = get_object_or_404(Account, id=account_id)
        
        # Get the new status from request
        new_status = request.POST.get('status')
        
        # Validate the status
        if new_status not in ['Active', 'Inactive']:
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid status provided'
            }, status=400)
        
        # Update the account status
        account.status = new_status
        account.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Account {account.user_id} has been {new_status.lower()}.',
            'new_status': new_status,
            'user_id': account.user_id
        })
    
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)
    
@csrf_exempt
@admin_required
def delete_account(request, account_id):
    if request.method == 'POST':
        acc = get_object_or_404(Account, id=account_id)
        acc.delete()
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'error'}, status=400)

@csrf_exempt
@admin_required
@require_POST
def update_account(request, account_id):
    try:
        data = json.loads(request.body)
        
        # Get the account
        account = Account.objects.get(id=account_id)
        
        # Update only the allowed fields
        account.role = data.get('role')
        account.email = data.get('email')
        
        # Handle course_section (assuming it's a ForeignKey)
        course_section_value = data.get('course_section')
        if course_section_value:
            account.course_section = CourseSection.objects.get(course_section=course_section_value)
        else:
            account.course_section = None
        
        account.save()
        
        return JsonResponse({
            'status': 'success',
            'message': f'Account {account.user_id} updated successfully'
        })
    except Account.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': 'Account not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=400)

import csv
import io
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from .models import Account, ClassSchedule, CourseSection, AttendanceRecord
from datetime import datetime

@csrf_exempt
@instructor_or_admin_required
@require_http_methods(["POST"])
def import_class_schedule(request):
    """Import attendance records from biometric CSV export - EXACT MATCH ONLY"""
    
    if 'csv_file' not in request.FILES:
        return JsonResponse({
            'status': 'error',
            'message': 'No CSV file uploaded',
            'imported': 0,
            'skipped': 0,
            'errors': ['No CSV file found']
        }, status=400)
    
    try:
        csv_file = request.FILES['csv_file']
        decoded_file = csv_file.read().decode('utf-8')
        io_string = io.StringIO(decoded_file)
        reader = csv.DictReader(io_string)
        
        results = {
            'attendance_created': 0,
            'attendance_updated': 0,
            'skipped': 0,
            'student_not_found': 0,
            'course_not_found': 0,
            'empty_rows': 0,
            'errors': []
        }
        
        for line_num, row in enumerate(reader, start=2):
            try:
                # ✨ EXTRACT FIELDS FROM YOUR CSV FORMAT
                course_code = row.get('course_code', '').strip()
                student_id = row.get('student_id', '').strip()
                attendance_date_str = row.get('attendance_date', '').strip()
                attendance_time_in_str = row.get('attendance_time_in', '').strip()
                attendance_time_out_str = row.get('attendance_time_out', '').strip()
                attendance_status = row.get('attendance_status', 'Present').strip().title()
                
                # ✨ SKIP COMPLETELY EMPTY ROWS
                if not any([course_code, student_id, attendance_date_str, attendance_time_in_str]):
                    results['empty_rows'] += 1
                    continue
                
                # ✨ VALIDATE REQUIRED FIELDS
                if not all([course_code, student_id, attendance_date_str, attendance_time_in_str]):
                    missing_fields = []
                    if not course_code: missing_fields.append('course_code')
                    if not student_id: missing_fields.append('student_id')
                    if not attendance_date_str: missing_fields.append('attendance_date')
                    if not attendance_time_in_str: missing_fields.append('attendance_time_in')
                    
                    results['errors'].append(f"Line {line_num}: Missing fields: {', '.join(missing_fields)}")
                    results['skipped'] += 1
                    continue
                
                # ✨ STEP 1: VERIFY COURSE CODE EXISTS
                try:
                    class_schedule = ClassSchedule.objects.get(course_code=course_code)
                except ClassSchedule.DoesNotExist:
                    results['course_not_found'] += 1
                    results['errors'].append(f"Line {line_num}: Course '{course_code}' not found")
                    results['skipped'] += 1
                    continue
                
                # ✨ STEP 2: VERIFY STUDENT EXISTS (EXACT MATCH ONLY - NO FLEXIBLE MATCHING)
                try:
                    student = Account.objects.get(user_id=student_id, role='Student')
                except Account.DoesNotExist:
                    results['student_not_found'] += 1
                    results['errors'].append(f"Line {line_num}: Student '{student_id}' not found (exact match required)")
                    results['skipped'] += 1
                    continue
                except Account.MultipleObjectsReturned:
                    results['student_not_found'] += 1
                    results['errors'].append(f"Line {line_num}: Multiple students found with ID '{student_id}'")
                    results['skipped'] += 1
                    continue
                
                # ✨ STEP 3: PARSE DATE & TIMES (FLEXIBLE FORMATS)
                try:
                    # Parse date (supports YYYY-MM-DD, DD/MM/YYYY, MM/DD/YYYY)
                    attendance_date = None
                    for date_format in ['%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y']:
                        try:
                            attendance_date = datetime.strptime(attendance_date_str, date_format).date()
                            break
                        except ValueError:
                            continue
                    
                    if not attendance_date:
                        raise ValueError(f"Date '{attendance_date_str}' must be YYYY-MM-DD, DD/MM/YYYY, or MM/DD/YYYY")
                    
                    # Parse time_in (supports H:M and H:M:S)
                    attendance_time_in = None
                    for time_format in ['%H:%M:%S', '%H:%M']:
                        try:
                            attendance_time_in = datetime.strptime(attendance_time_in_str, time_format).time()
                            break
                        except ValueError:
                            continue
                    
                    if not attendance_time_in:
                        raise ValueError(f"Time '{attendance_time_in_str}' must be HH:MM or HH:MM:SS")
                    
                    # Parse time_out (optional)
                    attendance_time_out = None
                    if attendance_time_out_str:
                        for time_format in ['%H:%M:%S', '%H:%M']:
                            try:
                                attendance_time_out = datetime.strptime(attendance_time_out_str, time_format).time()
                                break
                            except ValueError:
                                continue
                    
                except ValueError as e:
                    results['errors'].append(f"Line {line_num}: Invalid date/time - {str(e)}")
                    results['skipped'] += 1
                    continue
                
                # ✨ STEP 4: NORMALIZE STATUS
                status_map = {
                    'LATE': 'Late',
                    'PRESENT': 'Present',
                    'ABSENT': 'Absent',
                    'Late': 'Late',
                    'Present': 'Present',
                    'Absent': 'Absent'
                }
                normalized_status = status_map.get(attendance_status, 'Present')
                
                # ✨ STEP 5: CREATE OR UPDATE ATTENDANCE RECORD
                attendance_record, created = AttendanceRecord.objects.update_or_create(
                    class_schedule=class_schedule,
                    student=student,
                    date=attendance_date,
                    defaults={
                        'professor': class_schedule.professor,
                        'course_section': class_schedule.course_section,
                        'time_in': attendance_time_in,
                        'time_out': attendance_time_out,
                        'status': normalized_status,
                        'fingerprint_data': b'',
                    }
                )
                
                if created:
                    results['attendance_created'] += 1
                else:
                    results['attendance_updated'] += 1
                
            except Exception as e:
                results['errors'].append(f"Line {line_num}: {str(e)}")
                results['skipped'] += 1
        
        # ✨ DETERMINE STATUS
        total_processed = results['attendance_created'] + results['attendance_updated']
        
        if total_processed == 0:
            if results['skipped'] > 0:
                status_code = 'all_skipped'
            else:
                status_code = 'failed'
        elif results['skipped'] == 0:
            status_code = 'ok'
        else:
            status_code = 'partial'
        
        # ✨ BUILD SUMMARY MESSAGE
        message_parts = []
        if results['student_not_found'] > 0:
            message_parts.append(f"{results['student_not_found']} students not found")
        if results['course_not_found'] > 0:
            message_parts.append(f"{results['course_not_found']} courses not found")
        if results['empty_rows'] > 0:
            message_parts.append(f"{results['empty_rows']} empty rows skipped")
        
        summary_message = '; '.join(message_parts) if message_parts else 'All records processed successfully'
        
        print(f"✅ Import completed: {total_processed} records processed, {results['skipped']} skipped")
        
        return JsonResponse({
            'status': status_code,
            'imported': total_processed,
            'attendance_created': results['attendance_created'],
            'attendance_updated': results['attendance_updated'],
            'skipped': results['skipped'],
            'student_not_found': results['student_not_found'],
            'course_not_found': results['course_not_found'],
            'empty_rows': results['empty_rows'],
            'message': summary_message,
            'errors': results['errors'][:20]
        })
        
    except Exception as e:
        print(f"❌ Error in import: {str(e)}")
        import traceback
        traceback.print_exc()
        
        return JsonResponse({
            'status': 'error',
            'message': str(e),
            'imported': 0,
            'skipped': 0,
            'errors': [f'Server error: {str(e)}']
        }, status=500)
@require_http_methods(["POST"])
def import_class_from_pdf(request):
    """Import class schedule from PDF with duplicate detection and Unicode support"""
    if 'pdf_file' not in request.FILES:
        return JsonResponse({'status': 'error', 'message': 'No PDF file provided'}, status=400)
    
    pdf_file = request.FILES['pdf_file']
    
    try:
        # Read PDF content
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        full_text = ""
        
        # Extract text from all pages
        for page in pdf_reader.pages:
            full_text += page.extract_text() + "\n"
        
        # Clean up text - remove page breaks and normalize spacing
        full_text = re.sub(r'Page \d+ / \d+', '', full_text)
        full_text = re.sub(r'Page \d+ \d+', '', full_text)
        full_text = re.sub(r'\n{3,}', '\n\n', full_text)
        
        # Initialize data containers
        schedule_data = {}
        all_students = []
        
        # ========== SCHEDULE ID EXTRACTION ==========
        schedule_id_match = re.search(r'Schedule\s*ID\s*:\s*([A-Z0-9]+)', full_text, re.IGNORECASE)
        if schedule_id_match:
            schedule_data['schedule_id'] = schedule_id_match.group(1).strip()
        
        # ========== SUBJECT EXTRACTION ==========
        subject_match = re.search(r'Subject\s*:\s*([A-Z0-9-]+)\s*-\s*(.+?)\s+Venue', full_text, re.DOTALL)
        
        if not subject_match:
            subject_match = re.search(r'Subject\s*:\s*([A-Z0-9-]+)\s*-\s*(.+?)(?:\s+Course/Section|\n)', full_text, re.DOTALL)
        
        if not subject_match:
            subject_match = re.search(r'Subject\s*:\s*([A-Z0-9-]+)\s*-\s*(.+?)(?:\s{2,}|\n)', full_text)
        
        if subject_match:
            schedule_data['course_code'] = subject_match.group(1).strip()
            schedule_data['course_title'] = ' '.join(subject_match.group(2).strip().split())
        
        # ========== DAY/TIME EXTRACTION ==========
        day_time_match = re.search(r'Day/Time\s*:\s*([MTWRFSU])\s+(\d{1,2}:\d{2}\s*[AP]M)\s*-\s*(\d{1,2}:\d{2}\s*[AP]M)', full_text)
        
        if not day_time_match:
            day_time_match = re.search(r'Day/Time\s*:\s*([MTWRFSU]+)\s+(\d{1,2}:\d{2}\s*[AP]M)\s*-\s*(\d{1,2}:\d{2}\s*[AP]M)', full_text)
        
        if not day_time_match:
            day_time_match = re.search(r'Day/Time\s*:\s*([MTWRFSU/]+)\s+(\d{1,2}:\d{2}\s*[AP]M)\s*-\s*(\d{1,2}:\d{2}\s*[AP]M)', full_text)
        
        if day_time_match:
            day_map = {
                'M': 'Monday', 'T': 'Tuesday', 'W': 'Wednesday',
                'R': 'Thursday', 'F': 'Friday', 'S': 'Saturday', 'U': 'Sunday'
            }
            
            day_str = day_time_match.group(1).replace('/', '')
            if len(day_str) == 1:
                schedule_data['day'] = day_map.get(day_str, 'Monday')
            else:
                days = [day_map.get(d, d) for d in day_str if d in day_map]
                schedule_data['day'] = '/'.join(days)
            
            time_in_str = day_time_match.group(2).replace(' ', '')
            time_out_str = day_time_match.group(3).replace(' ', '')
            
            try:
                schedule_data['time_in'] = datetime.strptime(time_in_str, '%I:%M%p').time()
                schedule_data['time_out'] = datetime.strptime(time_out_str, '%I:%M%p').time()
            except ValueError:
                schedule_data['time_in'] = datetime.strptime(time_in_str.replace(' ', ''), '%I:%M%p').time()
                schedule_data['time_out'] = datetime.strptime(time_out_str.replace(' ', ''), '%I:%M%p').time()
        
        # ========== COURSE/SECTION EXTRACTION ==========
        section_match = re.search(r'Course/Section\s*:\s*(.+?)(?:\s*\n|1st Semester|2nd Semester|Summer)', full_text)
        
        if not section_match:
            section_match = re.search(r'Course/Section\s*:\s*(.+?)(?:\s+Student No\.|\n)', full_text)
        
        if section_match:
            section_str = section_match.group(1).strip()
            parts = section_str.split('-')
            
            if len(parts) >= 6:
                mid_point = len(parts) // 2
                schedule_data['course_name'] = '-'.join(parts[:mid_point])
                schedule_data['section_name'] = '-'.join(parts[-2:])
            elif len(parts) >= 3:
                mid_point = len(parts) // 2
                schedule_data['course_name'] = '-'.join(parts[:mid_point])
                schedule_data['section_name'] = '-'.join(parts[mid_point:])
            elif len(parts) == 2:
                schedule_data['course_name'] = parts[0]
                schedule_data['section_name'] = parts[1]
            else:
                schedule_data['course_name'] = section_str
                schedule_data['section_name'] = 'A'
        
        # ========== STUDENT EXTRACTION (ULTRA PERMISSIVE) ==========
        # Match anything that looks like: number + TUPC-ID + name + course code
        # Using .+? (any character) instead of specific character classes
        student_pattern = r'(\d+)\.+\s*(TUPC-\d{2}-\d{4})\s+(.+?)\s+([A-Z]{3,}[A-Z-]*)\s*(?:\n|$|Remarks|Page)'
        student_matches = list(re.finditer(student_pattern, full_text, re.DOTALL | re.UNICODE))
        
        print(f"DEBUG: Found {len(student_matches)} student matches")

        for match in student_matches:
            try:
                student_no = match.group(2).strip()
                name_raw = match.group(3).strip()
                
                # Clean the name - remove any non-letter characters except comma and space
                # This handles cases where special chars become weird symbols
                name = re.sub(r'[^\w\s,áéíóúÁÉÍÓÚñÑüÜ-]', '', name_raw)
                name = ' '.join(name.split())  # Normalize whitespace
                
                # Extract ID (TUPC-24-0107 → 240107)
                id_match = re.match(r'[A-Z]+-(\d{2})-(\d{4})', student_no)
                if not id_match:
                    print(f"DEBUG: Skipping - couldn't parse ID: {student_no}")
                    continue
                
                short_id = id_match.group(1) + id_match.group(2)
                
                # Parse name (LASTNAME, FIRSTNAME MIDDLENAME)
                if ',' not in name:
                    print(f"DEBUG: Skipping - no comma in name: {name}")
                    continue
                
                name_parts = name.split(',', 1)  # Split only on first comma
                if len(name_parts) < 2:
                    print(f"DEBUG: Skipping - couldn't split name: {name}")
                    continue
                
                last_name = name_parts[0].strip().title()
                first_name = name_parts[1].strip().title()  # Take ALL names after comma
                
                # Skip if either part is empty
                if not last_name or not first_name:
                    print(f"DEBUG: Skipping - empty name parts: '{last_name}' / '{first_name}'")
                    continue
                
                print(f"DEBUG: Parsed student #{match.group(1)} - ID: {short_id}, Name: {first_name} {last_name}")
                
                all_students.append({
                    'user_id': short_id,
                    'first_name': first_name,
                    'last_name': last_name
                })
                
            except Exception as e:
                print(f"DEBUG: Error parsing student entry: {str(e)}")
                import traceback
                traceback.print_exc()
                continue
        
        # ========== VALIDATION ==========
        required_fields = ['course_code', 'course_title', 'day', 'time_in', 'time_out', 'course_name', 'section_name']
        missing_fields = [f for f in required_fields if f not in schedule_data]
        
        if missing_fields:
            print(f"❌ Missing fields: {missing_fields}")
            print(f"DEBUG: Extracted schedule_data: {schedule_data}")
            return JsonResponse({
                'status': 'error',
                'message': f'Could not parse schedule information. Missing: {", ".join(missing_fields)}'
            }, status=400)
        
        if not all_students:
            print(f"❌ No students found in PDF")
            print(f"DEBUG: First 2000 chars of text:\n{full_text[:2000]}")
            return JsonResponse({
                'status': 'error',
                'message': 'Could not find any students in the PDF file'
            }, status=400)
        
        # ========== DATABASE OPERATIONS ==========
        course_section, created = CourseSection.objects.get_or_create(
            course_name=schedule_data['course_name'],
            section_name=schedule_data['section_name']
        )
        
        existing_schedule = ClassSchedule.objects.filter(
            course_section=course_section,
            days=schedule_data['day'],
            time_in=schedule_data['time_in'],
            time_out=schedule_data['time_out']
        ).first()
        
        if existing_schedule:
            return JsonResponse({
                'status': 'error',
                'message': f'Class schedule already exists for {course_section.course_section} on {schedule_data["day"]} from {schedule_data["time_in"]} to {schedule_data["time_out"]}'
            }, status=400)
        
        print(f"✅ Creating new ClassSchedule:")
        print(f"   Course/Section: {course_section.course_section}")
        print(f"   Students found: {len(all_students)}")
        
        class_schedule = ClassSchedule.objects.create(
            course_code=schedule_data['course_code'],
            course_title=schedule_data['course_title'],
            time_in=schedule_data['time_in'],
            time_out=schedule_data['time_out'],
            days=schedule_data['day'],
            course_section=course_section,
            professor=None,
            student_count=0,
            grace_period=15,
            remote_device='',
            room_assignment='-'
        )
        
        created_students = 0
        skipped_students = 0
        
        for student_info in all_students:
            if not Account.objects.filter(user_id=student_info['user_id']).exists():
                Account.objects.create(
                    user_id=student_info['user_id'],
                    email='',
                    first_name=student_info['first_name'],
                    last_name=student_info['last_name'],
                    role='Student',
                    password='00000',
                    sex='',
                    status='Pending',
                    course_section=course_section,
                    fingerprint_template=None
                )
                created_students += 1
            else:
                skipped_students += 1
        
        class_schedule.student_count = len(all_students)
        class_schedule.save()
        
        print(f"✅ Import completed: {created_students} students created, {skipped_students} skipped")
        
        return JsonResponse({
            'status': 'success',
            'message': 'Successfully imported class schedule',
            'details': {
                'course_code': schedule_data['course_code'],
                'course_title': schedule_data['course_title'],
                'course_section': course_section.course_section,
                'day': schedule_data['day'],
                'time': f"{schedule_data['time_in']} - {schedule_data['time_out']}",
                'students_created': created_students,
                'students_skipped': skipped_students,
                'total_students': len(all_students)
            }
        })
        
    except Exception as e:
        print(f"❌ Error importing PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({
            'status': 'error',
            'message': f'Failed to parse PDF: {str(e)}'
        }, status=500)



        
@admin_required
def class_management(request):
    
    # UPDATED: Get current active semester using is_active flag instead of date range
    current_semester = Semester.objects.filter(is_active=True).first()
    
    # Get unread notifications count
    new_accounts_count = AccountUploadNotification.objects.filter(is_read=False).count()
    recent_uploads = AccountUploadNotification.objects.filter(is_read=False)[:5]  # Last 5
    
    # Update student count
    schedules = ClassSchedule.objects.all()
    for schedule in schedules:
        student_acc = Account.objects.filter(course_section_id=schedule.course_section_id)
        student_count = len(student_acc)
        schedule.student_count = int(student_count)
        schedule.save()
    
    # Mark as read if user clicks "Mark as Read"
    if request.GET.get('mark_read') == 'true':
        AccountUploadNotification.objects.filter(is_read=False).update(is_read=True)
        messages.success(request, f'Marked {new_accounts_count} notifications as read.')
        return redirect('class_management')
    
    course_sections = CourseSection.objects.all()
    
    if request.method == 'POST':
        course_code = request.POST.get('course_code')
        course_name = request.POST.get('course_name')
        time_in = request.POST.get('time_in')
        time_out = request.POST.get('time_out')
        day = request.POST.get('day')
        course_section_str = request.POST.get('course_section')
        remote_device = request.POST.get('remote_device')
        
        try:
            section_obj = CourseSection.objects.get(course_section=course_section_str)
        except CourseSection.DoesNotExist:
            section_obj = None
        
        ClassSchedule.objects.create(
            course_code=course_code,
            course_title=course_name,
            time_in=time_in,
            time_out=time_out,
            days=day,
            course_section=section_obj,
            professor=None,
            student_count=0,
            grace_period=0,
            remote_device=remote_device,
            room_assignment='-',
        )
        return redirect('class_management')  # Redirect after POST to avoid resubmission
    
    # Pagination
    classes_list = ClassSchedule.objects.all().order_by('-id')
    paginator = Paginator(classes_list, 10)  # Show 10 classes per page
    
    page = request.GET.get('page', 1)
    try:
        classes = paginator.page(page)
    except PageNotAnInteger:
        classes = paginator.page(1)
    except EmptyPage:
        classes = paginator.page(paginator.num_pages)
    
    # Get instructors only
    instructors = Account.objects.filter(role="Instructor").values("id", "first_name", "last_name")
    instructors_json = json.dumps(list(instructors), cls=DjangoJSONEncoder)

    
    # UPDATED: Removed active_semester from context (not needed anymore)
    return render(request, 'class_management.html', {
        'course_sections': course_sections,
        'classes': classes,
        'instructors_json': instructors_json,
        'new_accounts_count': new_accounts_count,
        'recent_uploads': recent_uploads,
        'current_semester': current_semester,  # Only this is needed
    })


@require_http_methods(["POST"])
@admin_required
def add_course_section(request):
    """Add a new course section via AJAX"""
    print("🔵 add_course_section called")  # Debug log
    
    try:
        # Parse JSON body
        data = json.loads(request.body)
        course_name = data.get('course_name', '').strip()
        section_name = data.get('section_name', '').strip()

        print(f"📥 Received: course_name='{course_name}', section_name='{section_name}'")

        # Validate input
        if not course_name or not section_name:
            print("❌ Validation failed: missing fields")
            return JsonResponse({
                'status': 'error',
                'message': 'Course name and section name are required.'
            }, status=400)

        # Create course section string
        course_section_str = f"{course_name} {section_name}"
        
        # Check for duplicates
        if CourseSection.objects.filter(course_section=course_section_str).exists():
            print(f"⚠️ Duplicate detected: {course_section_str}")
            return JsonResponse({
                'status': 'error',
                'message': f'Section "{course_section_str}" already exists.'
            }, status=400)

        # Create new course section
        new_section = CourseSection.objects.create(
            course_name=course_name,
            section_name=section_name
        )
        
        print(f"✅ Created section: {new_section.course_section} (ID: {new_section.id})")

        # Return success response
        return JsonResponse({
            'status': 'success',
            'message': 'Course section added successfully.',
            'course_section': new_section.course_section
        })

    except json.JSONDecodeError as e:
        print(f"❌ JSON decode error: {e}")
        return JsonResponse({
            'status': 'error',
            'message': 'Invalid JSON format.'
        }, status=400)
    except Exception as e:
        print(f"❌ Exception in add_course_section: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)
@csrf_exempt
@admin_required
def update_class_schedule(request, pk):
    if request.method == "POST":
        try:
            print(f"🔵 update_class_schedule called for pk={pk}")
            
            cls = ClassSchedule.objects.get(id=pk)
            
            # Safer logging - don't rely on __str__ method
            print(f"✅ Found class schedule ID: {cls.id}, Course: {cls.course_code}")
            
            data = json.loads(request.body)
            print(f"📥 Received data: {data}")

            # Get professor by ID instead of parsing name
            prof_id = data.get("professor_id", "").strip()
            print(f"📋 Professor ID received: '{prof_id}'")
            
            if prof_id:
                try:
                    professor = Account.objects.get(id=prof_id, role="Instructor")
                    cls.professor = professor
                    print(f"✅ Assigned professor: {professor.first_name} {professor.last_name} (ID: {professor.id})")
                except Account.DoesNotExist:
                    cls.professor = None
                    print(f"⚠️ Professor with ID '{prof_id}' not found - setting to None")
            else:
                cls.professor = None
                print("ℹ️ No professor ID provided - setting to None")

            # Update other fields
            cls.time_in = data.get("time_in")
            cls.time_out = data.get("time_out")
            cls.days = data.get("day")
            cls.remote_device = data.get("remote_device")
            
            print(f"📝 Updating fields:")
            print(f"   time_in: {cls.time_in}")
            print(f"   time_out: {cls.time_out}")
            print(f"   days: {cls.days}")
            print(f"   remote_device: {cls.remote_device}")
            print(f"   professor: {cls.professor_id}")
            
            cls.save()
            print("✅ ClassSchedule saved successfully!")

            return JsonResponse({"status": "success"})
            
        except ClassSchedule.DoesNotExist:
            print(f"❌ ClassSchedule with pk={pk} does not exist")
            return JsonResponse({"status": "error", "message": "Class schedule not found"}, status=404)
            
        except json.JSONDecodeError as e:
            print(f"❌ JSON decode error: {e}")
            return JsonResponse({"status": "error", "message": "Invalid JSON"}, status=400)
            
        except Exception as e:
            print(f"❌ Unexpected exception: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"status": "error", "message": str(e)}, status=500)
    
    print(f"⚠️ Invalid request method: {request.method}")
    return JsonResponse({"status": "error", "message": "Invalid request method"}, status=400)




@csrf_exempt
@admin_required
def delete_class_schedule(request, pk):
    if request.method == "POST":
        try:
            cls = ClassSchedule.objects.get(id=pk)
            cls.delete()
            return JsonResponse({"status": "deleted"})
        except:
            return JsonResponse({"status": "error"}, status=400)

@instructor_or_admin_required
def attendance_report_template(request):
    # Get logged-in instructor's Account entry
    try:
        instructor_account = Account.objects.get(email=request.user.email, role='Instructor')
    except Account.DoesNotExist:
        return render(request, "error.html", {"message": "Instructor account not found"})
    
    # Get class schedules for the dropdown
    schedules = ClassSchedule.objects.filter(professor=instructor_account)
    
    selected_schedule_id = request.GET.get("schedule")
    selected_date_range = request.GET.get("date_range")
    date_ranges = []
    attendance_table = []
    attendance_data = {}
    students_list = []
    
    if selected_schedule_id:
        try:
            # Get the selected schedule
            schedule_obj = ClassSchedule.objects.get(id=selected_schedule_id)
            
            # Class details for the form
            attendance_data_attendance_report = {
                'subject': schedule_obj.course_title,
                'faculty_name': f"{schedule_obj.professor.first_name} {schedule_obj.professor.last_name}" if schedule_obj.professor else 'TBA',
                'course': schedule_obj.course_section.course_name if schedule_obj.course_section else '',
                'room': schedule_obj.room_assignment or 'TBA',
                'year_section': schedule_obj.course_section.section_name if schedule_obj.course_section else '',
                'schedule': f"{schedule_obj.days} {schedule_obj.time_in}-{schedule_obj.time_out}",
            }
            
        except ClassSchedule.DoesNotExist:
            pass

        # Fetch unique attendance dates for selected class schedule
        attendance_dates = AttendanceRecord.objects.filter(
            class_schedule_id=selected_schedule_id
        ).values_list("date", flat=True).distinct().order_by("date")

        attendance_dates = list(attendance_dates)

        # Group into 8-day ranges for the filter dropdown
        for i in range(0, len(attendance_dates), 8):
            start_date = attendance_dates[i]
            end_date = attendance_dates[min(i + 7, len(attendance_dates) - 1)]
            date_ranges.append({
                "value": f"{start_date}_to_{end_date}",
                "label": f"{start_date.strftime('%B %d, %Y')} - {end_date.strftime('%B %d, %Y')}"
            })

        #if may sinelect na na date range
        if selected_date_range:
            try:
                start_str, end_str = selected_date_range.split("_to_")
                start_date = parse_date(start_str)
                end_date = parse_date(end_str)
            except (ValueError, TypeError):
                start_date = end_date = None

            if start_date and end_date:
                # Get all attendance records within the date range
                attendance_qs = AttendanceRecord.objects.filter(
                    class_schedule_id=selected_schedule_id,
                    date__range=(start_date, end_date)
                ).select_related('student')

                # Get schedule object once
                schedule_obj = ClassSchedule.objects.get(id=selected_schedule_id)

                # Get students in the same course_section as the schedule
                students_in_schedule = Account.objects.filter(
                    course_section_id=schedule_obj.course_section_id
                ).order_by('last_name', 'first_name')

                # Map attendance data: {student_id: {date: {'status': status, 'time_in': time_in, 'time_out': time_out}}}
                attendance_data = defaultdict(lambda: defaultdict(dict))
                for record in attendance_qs:
                    attendance_data[record.student.id][record.date] = {
                        'status': record.status,
                        'time_in': record.time_in,
                        'time_out': record.time_out
                    }

                # Build date headers (max 8 dates)
                date_headers = [d for d in attendance_dates if start_date <= d <= end_date][:8]
                num_empty_date_column = 8 - len(date_headers)
                #create a pseudo list just for the for loop in html to work
                num_empty_date_columns = []
                for i in range(num_empty_date_column):
                    num_empty_date_columns.append("")


                attendance_table = []
                for student in students_in_schedule:
                    course_section_for_student = student.course_section
                    # Build dates_statuses as a list of dicts containing all attendance info
                    dates_statuses = []
                    for date in date_headers:
                        attendance_info = attendance_data[student.id].get(date, {})
                        dates_statuses.append({
                            'status': attendance_info.get('status', ''),
                            'time_in': attendance_info.get('time_in', ''),
                            'time_out': attendance_info.get('time_out', '')
                        })
                    
                    # Pad with empty dicts to always have 8 items
                    while len(dates_statuses) < 8:
                        dates_statuses.append({
                            'status': '',
                            'time_in': '',
                            'time_out': ''
                        })
                    
                        
                    row = {
                        "student_id": student.user_id, 
                        "name": f"{student.first_name} {student.last_name}",
                        "sex": student.sex,
                        "course": course_section_for_student,
                        "subject": schedule_obj.course_code,
                        "room": schedule_obj.room_assignment,
                        "dates": dates_statuses,  # Always 8 items now
                    }
                    attendance_table.append(row)

                else:
                    row = [""]
                # Build date headers (max 8 dates)
                date_headers = [d for d in attendance_dates if start_date <= d <= end_date][:8]
                num_empty_date_column = 8 - len(date_headers)
                #create a pseudo list just for the for loop in html to work
                num_empty_date_columns = []
                for i in range(num_empty_date_column):
                    num_empty_date_columns.append("")

                num_empty_row = 40 - len(attendance_table)
                #create a pseudo list just for the for loop in html to work
                num_empty_rows = []
                num_filled_rows = []
                for i in range(num_empty_row):
                    num_empty_rows.append("")

                for i in range(len(attendance_table)):
                    num_filled_rows.append("")

                print(f"{len(num_empty_rows)}")
                print(f"{len(num_filled_rows)}")

                context = {
                    'attendance_data': attendance_data_attendance_report,
                    "schedules": schedules,
                    "date_ranges": date_ranges,
                    "selected_schedule_id": selected_schedule_id,
                    "selected_date_range": selected_date_range,
                    "attendance_table": attendance_table,
                    "date_headers": date_headers,
                    "num_empty_rows": num_empty_rows,
                    "num_filled_rows": num_filled_rows,
                    "num_empty_date_columns": num_empty_date_columns,
                }
                return render(request, "attendance_report_template.html", context)
        else:
    # If no valid date range selected or dates invalid, still render with schedules & date ranges
            context = {
                'attendance_data': attendance_data_attendance_report,
                "schedules": schedules,
                "date_ranges": date_ranges,
                "selected_schedule_id": selected_schedule_id,
                "attendance_table": [],
                "date_headers": [],
                "num_empty_rows": [""]*40,
                "num_filled_rows": [],
                "num_empty_date_columns": ["", "", "", "", "", "", "", "", ],
            }
            return render(request, "attendance_report_template.html", context)

    # If no schedule selected at all, render with just schedules and empty date ranges
    context = {
        'attendance_data': {"":"",'faculty_name': f"{instructor_account.first_name} {instructor_account.last_name}" if instructor_account else 'TBA',"":"","":"","":"","":""},
        "schedules": schedules,
        "date_ranges": date_ranges,
        "selected_schedule_id": selected_schedule_id,
        "attendance_table": [],
        "date_headers": [],
        "num_empty_rows": [""]*40,
        "num_filled_rows": [],
        "num_empty_date_columns": ["", "", "", "", "", "", "", "", ],
    }
    return render(request, "attendance_report_template.html", context)
    

from django.http import JsonResponse
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime, date
from django.utils import timezone

@admin_required
@require_POST
def set_semester(request):
    """Set a new semester - can only be done once until end date passes"""
    try:
        semester_name = request.POST.get('semester_name')
        school_year = request.POST.get('school_year')
        start_date_str = request.POST.get('start_date')
        end_date_str = request.POST.get('end_date')
        
        # Validation - Check all required fields
        if not all([semester_name, school_year, start_date_str, end_date_str]):
            return JsonResponse({
                'status': 'error',
                'message': 'All fields are required'
            }, status=400)
        
        # Parse dates
        try:
            start_date = datetime.strptime(start_date_str, '%Y-%m-%d').date()
            end_date = datetime.strptime(end_date_str, '%Y-%m-%d').date()
        except ValueError:
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid date format. Use YYYY-MM-DD'
            }, status=400)
        
        # Validate end date is after start date
        if end_date <= start_date:
            return JsonResponse({
                'status': 'error',
                'message': 'End date must be after start date'
            }, status=400)
        
        # Get current active semester
        today = date.today()
        current_semester = Semester.objects.filter(is_active=True).first()
        
        # Check if there's an active semester that hasn't ended yet
        if current_semester:
            if current_semester.end_date >= today:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Cannot set new semester. Current semester "{current_semester.semester_name} ({current_semester.school_year})" is still active until {current_semester.end_date.strftime("%B %d, %Y")}. You can only set a new semester after this date.'
                }, status=400)
        
        # Validate start date is not in the past (optional - remove if you want to allow past dates)
        if start_date < today:
            return JsonResponse({
                'status': 'error',
                'message': 'Semester start date cannot be earlier than today'
            }, status=400)
        
        # Check for duplicate semester
        if Semester.objects.filter(
            semester_name=semester_name,
            school_year=school_year
        ).exists():
            return JsonResponse({
                'status': 'error',
                'message': f'{semester_name} for school year {school_year} already exists'
            }, status=400)
        
        # Archive previous semester if it exists and has ended
        if current_semester:
            current_semester.is_active = False
            current_semester.is_archived = True
            current_semester.save()
        
        # Create new semester
        new_semester = Semester.objects.create(
            semester_name=semester_name,
            school_year=school_year,
            start_date=start_date,
            end_date=end_date,
            is_active=True,
            is_archived=False
        )
        
        return JsonResponse({
            'status': 'success',
            'message': 'Semester set successfully',
            'semester': {
                'id': new_semester.id,
                'semester_name': new_semester.semester_name,
                'school_year': new_semester.school_year,
                'start_date': new_semester.start_date.strftime('%B %d, %Y'),
                'end_date': new_semester.end_date.strftime('%B %d, %Y')
            }
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Server error: {str(e)}'
        }, status=500)


# API views for mobile app integration
class AccountViewSet(viewsets.ModelViewSet):
    queryset = Account.objects.all()
    serializer_class = AccountSerializer
    
    def get_permissions(self):
        """Require authentication for write operations"""
        if self.action in ['list', 'retrieve']:
            return [AllowAny()]
        return [IsAuthenticated()]
    
    def create(self, request, *args, **kwargs):
        """Override create to add debugging"""
        print("=== DEBUG: Incoming account data ===")
        print("Request data:", request.data)
        print("Request headers:", dict(request.headers))
        
        serializer = self.get_serializer(data=request.data)
        if not serializer.is_valid():
            print("=== DEBUG: Validation errors ===")
            print("Serializer errors:", serializer.errors)
            return Response(serializer.errors, status=400)
        
        return super().create(request, *args, **kwargs) 
    
    def perform_create(self, serializer):
        """Add audit trail and notification for account creation"""
        account = serializer.save()
        
        # Create notification for new account upload
        AccountUploadNotification.objects.create(
            account_name=f"{account.first_name} {account.last_name}"
        )
        
        # Log who created this account
        if self.request.user.username == 'mobile_system':
            print(f"Account {account.user_id} created via mobile upload")
            print(f"Notification created for new account: {account.first_name} {account.last_name}")
        else:
            print(f"Account {account.user_id} created by user {self.request.user.username}")
        
    @action(detail=False, methods=['post'], permission_classes=[IsAuthenticated])
    def mobile_sync(self, request):
        """Sync accounts from mobile to web"""
        serializer = MobileAccountSerializer(data=request.data, many=True)
        if serializer.is_valid():
            return Response({'status': 'success'})
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# Web app to mobile app account overwrite/syncing
@csrf_exempt
@require_http_methods(["GET"])
def mobile_account_sync(request):
    """API endpoint for mobile apps to fetch all accounts"""
    try:
        # Get all accounts from Django database
        accounts = Account.objects.all()
        
        # Convert to list and format for mobile consumption
        accounts_list = []
        for account in accounts:
            accounts_list.append({
                'user_id': account.user_id,
                'email': account.email,
                'first_name': account.first_name,
                'last_name': account.last_name,
                'role': account.role,
                'password': None,  # Don't send passwords
                'sex': account.sex,
                'status': account.status,
                'course_section': account.course_section.id if account.course_section else None,
                'fingerprint_template': account.fingerprint_template
            })
        
        return JsonResponse(accounts_list, safe=False)
        
    except Exception as e:
        return JsonResponse({
            'error': str(e)
        }, status=500)
        
@csrf_exempt
@require_http_methods(["POST"])
def trigger_mobile_sync(request):
    """Trigger sync to mobile apps"""
    try:
        print("Mobile sync triggered from web admin")
        
        # Get counts of data available for sync
        account_count = Account.objects.count()
        schedule_count = ClassSchedule.objects.count()
        course_section_count = CourseSection.objects.count()
        
        return JsonResponse({
            'status': 'success',
            'message': 'Mobile sync triggered successfully',
            'data': {
                'accounts_available': account_count,
                'schedules_available': schedule_count,
                'course_sections_available': course_section_count
            }
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

class ClassScheduleViewSet(viewsets.ModelViewSet):
    queryset = ClassSchedule.objects.all()
    serializer_class = MobileClassScheduleSerializer
    
    def get_permissions(self):
        """Allow read without auth, require auth for write operations"""
        if self.action in ['list', 'retrieve', 'today_schedules']:
            return [AllowAny()]
        return [IsAuthenticated()]

    @action(detail=False, methods=['get'])
    def today_schedules(self, request):
        """Get today's schedules for mobile"""
        today = timezone.now().date()
        schedules = ClassSchedule.objects.all()
        # Make sure this line uses MobileClassScheduleSerializer
        serializer = MobileClassScheduleSerializer(schedules, many=True)
        return Response(serializer.data)

class AttendanceRecordViewSet(viewsets.ModelViewSet):
    queryset = AttendanceRecord.objects.all()
    serializer_class = AttendanceRecordSerializer
    
    def get_permissions(self):
        """Allow read without auth, require auth for write operations"""
        if self.action in ['list', 'retrieve', 'download_for_mobile']:
            return [AllowAny()]
        return [IsAuthenticated()]

    @action(detail=False, methods=['post'], permission_classes=[IsAuthenticated])
    def mobile_upload(self, request):
        """Upload attendance records from mobile"""
        # Explicitly use MobileAttendanceSerializer (which has the user_id lookup logic)
        serializer = MobileAttendanceSerializer(data=request.data, many=True)
        if serializer.is_valid():
            created_records = serializer.save()
            return Response({
                'status': 'success', 
                'count': len(created_records),
                'message': f'Successfully uploaded {len(created_records)} attendance records'
            })
        else:
            print("Validation errors:", serializer.errors)  # Debug print
            return Response({
                'status': 'error',
                'errors': serializer.errors
            }, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['get'])
    def download_for_mobile(self, request):
        """Download attendance records to mobile"""
        date_param = request.query_params.get('date')
        if date_param:
            try:
                filter_date = datetime.strptime(date_param, '%Y-%m-%d').date()
                records = AttendanceRecord.objects.filter(date=filter_date)
            except ValueError:
                return Response({'error': 'Invalid date format'}, status=status.HTTP_400_BAD_REQUEST)
        else:
            records = AttendanceRecord.objects.all()
        
        serializer = MobileAttendanceSerializer(records, many=True)
        return Response(serializer.data)


class CourseSectionViewSet(viewsets.ModelViewSet):
    """
    API endpoint for course sections - allows mobile app to fetch all course sections
    """
    queryset = CourseSection.objects.all()
    serializer_class = CourseSectionSerializer
    
    def get_permissions(self):
        """Allow read without auth for mobile sync"""
        if self.action in ['list', 'retrieve']:
            return [AllowAny()]
        return [IsAuthenticated()]
    
@api_view(['POST'])
@permission_classes([AllowAny])
def mobile_login(request):
    """Login endpoint for mobile app"""
    username = request.data.get('username')
    password = request.data.get('password')
    
    if username and password:
        user = authenticate(username=username, password=password)
        if user:
            token, created = Token.objects.get_or_create(user=user)
            return Response({
                'token': token.key,
                'user_id': user.id,
                'username': user.username
            })
    
    return Response({'error': 'Invalid credentials'}, status=status.HTTP_401_UNAUTHORIZED)


@api_view(['POST'])
@permission_classes([AllowAny])
def mobile_auth(request):
    """Authentication endpoint for mobile devices"""
    device_id = request.data.get('device_id')
    device_secret = request.data.get('device_secret')
    
    # You can store these in Django settings or environment variables
    VALID_DEVICES = {
        'FINGERPRINT_DEVICE_001': 'Room1_Debug_2025', 
        # Add more devices as needed
    }
    
    if device_id in VALID_DEVICES and VALID_DEVICES[device_id] == device_secret:
        # Create or get a system user for mobile uploads
        user, created = User.objects.get_or_create(
            username='mobile_system',
            defaults={'email': 'mobile@system.local', 'is_active': True}
        )
        
        # Create or get token for this user
        token, created = Token.objects.get_or_create(user=user)
        
        return Response({
            'token': token.key,
            'expires_in': 86400,  # 24 hours
            'device_id': device_id
        })
    
    return Response({'error': 'Invalid device credentials'}, status=401)
@csrf_exempt
@require_http_methods(["PUT", "PATCH", "POST"])
def mobile_update_account(request, user_id):
    """Mobile-specific account update endpoint - allows devices to complete pending accounts"""
    try:
        # Optional: Basic device authentication check
        device_id = request.headers.get('X-Device-ID')
        
        # Get the account
        try:
            account = Account.objects.get(user_id=user_id)
        except Account.DoesNotExist:
            return JsonResponse({'error': 'Account not found'}, status=404)
        
        # Parse update data
        data = json.loads(request.body)
        
        # Update only allowed fields (email, sex, fingerprint, status)
        if 'email' in data and data['email']:
            account.email = data['email']
        if 'sex' in data and data['sex']:
            account.sex = data['sex']
        if 'fingerprint_template' in data:
            account.fingerprint_template = data['fingerprint_template']
        if 'status' in data:
            account.status = data['status']
            
        account.save()
        
        if was_pending and account.status == 'Active':
            AccountUploadNotification.objects.create(
                account_name=f"{account.first_name} {account.last_name}",
                notification_type='update'
            )
            
        # Return updated account
        return JsonResponse({
            'user_id': account.user_id,
            'email': account.email,
            'first_name': account.first_name,
            'last_name': account.last_name,
            'role': account.role,
            'password': None,
            'sex': account.sex,
            'status': account.status,
            'course_section': account.course_section.id if account.course_section else None,
            'fingerprint_template': account.fingerprint_template
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@instructor_or_admin_required
@require_http_methods(["GET"])
def get_attendance_data_api(request):
    """API endpoint to get attendance data as JSON"""
    
    try:
        schedule_id = request.GET.get('schedule')
        date_range = request.GET.get('date_range')
        
        if not schedule_id:
            return JsonResponse({'error': 'No schedule selected'}, status=400)
        
        try:
            class_schedule = ClassSchedule.objects.get(id=schedule_id)
        except ClassSchedule.DoesNotExist:
            return JsonResponse({'error': 'Class schedule not found'}, status=404)
        
        # Get all date ranges
        attendance_dates = list(AttendanceRecord.objects.filter(
            class_schedule=class_schedule
        ).values_list('date', flat=True).distinct().order_by('date'))
        
        date_ranges = []
        for i in range(0, len(attendance_dates), 8):
            start_date = attendance_dates[i]
            end_date = attendance_dates[min(i + 7, len(attendance_dates) - 1)]
            date_ranges.append({
                'value': f'{start_date}_to_{end_date}',
                'label': f'{start_date.strftime("%b %d, %Y")} - {end_date.strftime("%b %d, %Y")}'
            })
        
        # If date range provided, build attendance table
        if date_range:
            try:
                start_str, end_str = date_range.split('_to_')
                start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
            except (ValueError, TypeError):
                start_date = end_date = None
        else:
            start_date = end_date = None
        
        # Get date headers
        if start_date and end_date:
            dates_in_range = list(AttendanceRecord.objects.filter(
                class_schedule=class_schedule,
                date__range=[start_date, end_date]
            ).values_list('date', flat=True).distinct().order_by('date')[:8])
        else:
            dates_in_range = attendance_dates[:8]
        
        date_headers = [d.strftime('%m/%d') for d in dates_in_range]
        
        # Get students and attendance
        attendance_table = []
        students = Account.objects.filter(
            course_section=class_schedule.course_section,
            role='Student'
        ).order_by('last_name', 'first_name')
        
        # Build attendance data map
        attendance_data = defaultdict(lambda: defaultdict(dict))
        attendance_qs = AttendanceRecord.objects.filter(
            class_schedule=class_schedule
        ).select_related('student')
        
        for record in attendance_qs:
            attendance_data[record.student.id][record.date] = {
                'time_in': record.time_in,
                'time_out': record.time_out,
                'status': record.status
            }
        
        # Build attendance table
        for student in students:
            student_data = {
                'name': f'{student.last_name}, {student.first_name}',
                'sex': student.sex[0] if student.sex else 'M',
                'dates': []
            }
            
            for date in dates_in_range:
                if date in attendance_data[student.id]:
                    att = attendance_data[student.id][date]
                    student_data['dates'].append({
                        'time_in': att['time_in'].strftime('%H:%M') if att['time_in'] else '',
                        'time_out': att['time_out'].strftime('%H:%M') if att['time_out'] else '',
                        'status': att['status']
                    })
                else:
                    student_data['dates'].append({
                        'time_in': '',
                        'time_out': '',
                        'status': ''
                    })
            
            attendance_table.append(student_data)
        
        # Get class data
        instructor_account = class_schedule.professor
        class_data = {
            'subject': class_schedule.course_title,
            'faculty_name': f'{instructor_account.first_name} {instructor_account.last_name}' if instructor_account else 'TBA',
            'course': class_schedule.course_code,
            'room': class_schedule.room_assignment,
            'year_section': class_schedule.course_section.section_code if class_schedule.course_section else 'N/A',
            'schedule': f'{class_schedule.days} {class_schedule.time_in.strftime("%H:%M")}-{class_schedule.time_out.strftime("%H:%M")}'
        }
        
        return JsonResponse({
            'class_data': class_data,
            'date_ranges': date_ranges,
            'date_headers': date_headers,
            'attendance_table': attendance_table,
            'student_count': len(attendance_table)
        })
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"ERROR in get_attendance_data_api: {error_trace}")
        return JsonResponse({
            'error': f'Server error: {str(e)}',
            'trace': error_trace
        }, status=500)


# for Docx Download
from io import BytesIO
from docx import Document
from collections import defaultdict
import os
import re
from datetime import datetime
import logging
from PyPDF2 import PdfMerger
import subprocess
import tempfile




@instructor_or_admin_required
def generate_attendance_docx(request, schedule_id):
    """Generate Attendance Report - 60 Students, Both Templates, DOCX Output"""
    logger = logging.getLogger(__name__)
    logger.error(f"=== DOCX Download Started for schedule_id: {schedule_id} ===")
    
    date_range = request.GET.get('date_range')
    if not date_range:
        logger.error("✗ No date range provided")
        return HttpResponse('<h3>Date Range Required</h3><p>Please select a date range.</p>', status=400)
    
    try:
        logger.error(f"Raw date_range: '{date_range}'")
        parts = date_range.split('to')
        start_str = re.sub(r'[^0-9-]', '', parts[0].strip())
        end_str = re.sub(r'[^0-9-]', '', parts[1].strip())
        start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
        date_range_str = f"_{start_date.strftime('%m%d')}-{end_date.strftime('%m%d')}"
        logger.error(f"✓ Parsed: {start_date} to {end_date}")
    except Exception as e:
        logger.error(f"✗ Invalid date: {str(e)}")
        return HttpResponse(f'<h3>Invalid Date Range</h3><p>{str(e)}</p>', status=400)

    # Load BOTH templates
    template1_path = os.path.join(settings.BASE_DIR, 'PTLT_App', 'templates', 'attendance_template.docx')
    template2_path = os.path.join(settings.BASE_DIR, 'PTLT_App', 'templates', 'attendance_template2.docx')
    
    if not os.path.exists(template1_path) or not os.path.exists(template2_path):
        return HttpResponse("Templates not found", status=500)
    
    logger.error("✓ Both templates loaded")

    class_schedule = ClassSchedule.objects.get(id=schedule_id)
    
    # Get up to 60 students
    students_qs = Account.objects.filter(
        course_section=class_schedule.course_section,
        role='Student'
    ).order_by('last_name', 'first_name')

    seen_names = set()
    students = []
    for student in students_qs:
        full_name = f"{student.last_name},{student.first_name}"
        if full_name not in seen_names:
            seen_names.add(full_name)
            students.append(student)
            if len(students) >= 60:
                break
    
    logger.error(f"✓ {len(students)} students")

    attendance_dates = list(AttendanceRecord.objects.filter(
        class_schedule=class_schedule,
        date__range=[start_date, end_date]
    ).values_list('date', flat=True).distinct().order_by('date')[:8])
    
    attendance_qs = AttendanceRecord.objects.filter(
        class_schedule=class_schedule,
        date__range=[start_date, end_date]
    ).select_related('student')

    # Map attendance data
    attendance_data = defaultdict(lambda: defaultdict(dict))
    for record in attendance_qs:
        attendance_data[record.student.id][record.date] = {
            'time_in': record.time_in,
            'time_out': record.time_out,
            'status': record.status,
            'professor_time_in': record.professor_time_in,
            'professor_time_out': record.professor_time_out
        }
    logger.error(f"✓ {len(attendance_dates)} dates")

    # Split students: 40 for Template1, 20 for Template2
    students_template1 = students[0:40]
    students_template2 = students[40:60]
    logger.error(f"✓ Template1: {len(students_template1)}, Template2: {len(students_template2)}")

    # ==================== TEMPLATE 1 PROCESSING ====================
    doc1 = Document(template1_path)
    
    replacements1 = {
        '{{subject}}': class_schedule.course_title or class_schedule.course_code,
        '{{faculty_name}}': f"{class_schedule.professor.first_name} {class_schedule.professor.last_name}" if class_schedule.professor else "TBA",
        '{{course}}': class_schedule.course_section.course_name if class_schedule.course_section else "",
        '{{room_assignment}}': class_schedule.room_assignment or "TBA",
        '{{year_section}}': class_schedule.course_section.section_name if class_schedule.course_section else "",
        '{{schedule}}': f"{class_schedule.days} {class_schedule.time_in.strftime('%H:%M')}-{class_schedule.time_out.strftime('%H:%M')}"
    }

    # Date headers with professor time
    for i in range(1, 9):
        if i - 1 < len(attendance_dates):
            date_obj = attendance_dates[i-1]
            date_str = date_obj.strftime('%m/%d/%Y')
            att_for_date = AttendanceRecord.objects.filter(
                class_schedule=class_schedule,
                date=date_obj
            ).exclude(professor_time_in=None, professor_time_out=None).first()
            if att_for_date and att_for_date.professor_time_in and att_for_date.professor_time_out:
                prof_time = f"\n{att_for_date.professor_time_in.strftime('%H:%M')}-{att_for_date.professor_time_out.strftime('%H:%M')}"
                replacements1[f'{{{{date{i}}}}}'] = date_str + prof_time
            else:
                replacements1[f'{{{{date{i}}}}}'] = date_str
        else:
            replacements1[f'{{{{date{i}}}}}'] = ''

    # Student data for Template1
    time_cells1 = set()
    for i in range(1, 41):
        if i - 1 < len(students_template1):
            student = students_template1[i - 1]
            replacements1[f'{{{{student{i}_name}}}}'] = f"{student.last_name}, {student.first_name}"
            replacements1[f'{{{{student{i}_sex}}}}'] = student.sex[0] if student.sex else ''
            
            for j in range(1, 9):
                key = f'{{{{student{i}_time{j}}}}}'
                if j - 1 < len(attendance_dates):
                    date = attendance_dates[j - 1]
                    if date in attendance_data[student.id]:
                        att = attendance_data[student.id][date]
                        if att['status'] in ['Present', 'Late']:
                            time_in_str = att['time_in'].strftime('%H:%M') if att['time_in'] else ''
                            time_out_str = att['time_out'].strftime('%H:%M') if att['time_out'] else ''
                            if time_in_str and time_out_str:
                                replacements1[key] = f"{time_in_str} - {time_out_str}"
                                time_cells1.add(key)
                                continue
                replacements1[key] = ''
        else:
            replacements1[f'{{{{student{i}_name}}}}'] = ''
            replacements1[f'{{{{student{i}_sex}}}}'] = ''
            for j in range(1, 9):
                replacements1[f'{{{{student{i}_time{j}}}}}'] = ''

    logger.error(f"✓ Built {len(replacements1)} replacements for Template1")

    # Apply replacements to Template1 WITH FONT STYLING
    for paragraph in doc1.paragraphs:
        for key, value in replacements1.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))

    for table in doc1.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements1.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value))
                                    if key in time_cells1:
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(8)

    # ==================== TEMPLATE 2 PROCESSING (still filled, but not returned here) ====================
    doc2 = Document(template2_path)
    
    replacements2 = {
        '{{subject}}': class_schedule.course_title or class_schedule.course_code,
        '{{faculty_name}}': f"{class_schedule.professor.first_name} {class_schedule.professor.last_name}" if class_schedule.professor else "TBA",
        '{{course}}': class_schedule.course_section.course_name if class_schedule.course_section else "",
        '{{room_assignment}}': class_schedule.room_assignment or "TBA",
        '{{year_section}}': class_schedule.course_section.section_name if class_schedule.course_section else "",
        '{{schedule}}': f"{class_schedule.days} {class_schedule.time_in.strftime('%H:%M')}-{class_schedule.time_out.strftime('%H:%M')}"
    }

    for i in range(1, 9):
        if i - 1 < len(attendance_dates):
            date_obj = attendance_dates[i-1]
            date_str = date_obj.strftime('%m/%d/%Y')
            att_for_date = AttendanceRecord.objects.filter(
                class_schedule=class_schedule,
                date=date_obj
            ).exclude(professor_time_in=None, professor_time_out=None).first()
            if att_for_date and att_for_date.professor_time_in and att_for_date.professor_time_out:
                prof_time = f"\n{att_for_date.professor_time_in.strftime('%H:%M')}-{att_for_date.professor_time_out.strftime('%H:%M')}"
                replacements2[f'{{{{date{i}}}}}'] = date_str + prof_time
            else:
                replacements2[f'{{{{date{i}}}}}'] = date_str
        else:
            replacements2[f'{{{{date{i}}}}}'] = ''

    time_cells2 = set()
    for i in range(1, 41):
        if i - 1 < len(students_template2):
            student = students_template2[i - 1]
            replacements2[f'{{{{student{i}_name}}}}'] = f"{student.last_name}, {student.first_name}"
            replacements2[f'{{{{student{i}_sex}}}}'] = student.sex[0] if student.sex else ''
            
            for j in range(1, 9):
                key = f'{{{{student{i}_time{j}}}}}'
                if j - 1 < len(attendance_dates):
                    date = attendance_dates[j - 1]
                    if date in attendance_data[student.id]:
                        att = attendance_data[student.id][date]
                        if att['status'] in ['Present', 'Late']:
                            time_in_str = att['time_in'].strftime('%H:%M') if att['time_in'] else ''
                            time_out_str = att['time_out'].strftime('%H:%M') if att['time_out'] else ''
                            if time_in_str and time_out_str:
                                replacements2[key] = f"{time_in_str} - {time_out_str}"
                                time_cells2.add(key)
                                continue
                replacements2[key] = ''
        else:
            replacements2[f'{{{{student{i}_name}}}}'] = ''
            replacements2[f'{{{{student{i}_sex}}}}'] = ''
            for j in range(1, 9):
                replacements2[f'{{{{student{i}_time{j}}}}}'] = ''

    logger.error(f"✓ Built {len(replacements2)} replacements for Template2")

    for paragraph in doc2.paragraphs:
        for key, value in replacements2.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))

    for table in doc2.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements2.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value))
                                    if key in time_cells2:
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(8)

    # ==================== RETURN DOCX (Template 1) ====================
    out = BytesIO()
    doc1.save(out)
    out.seek(0)

    filename = f"attendance_{schedule_id}{date_range_str}.docx"
    response = HttpResponse(
        out.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    logger.error("✓ DOCX sent to user")
    return response



# for pdf preview also

@instructor_or_admin_required
def get_attendance_data_api(request):
    """API endpoint to get attendance data as JSON"""
    schedule_id = request.GET.get('schedule')
    date_range = request.GET.get('date_range')
    
    if not schedule_id:
        return JsonResponse({'error': 'No schedule selected'}, status=400)
    
    try:
        class_schedule = ClassSchedule.objects.get(id=schedule_id)
        
        # Get class details
        class_data = {
            'subject': class_schedule.course_title or class_schedule.course_code,
            'faculty_name': f"{class_schedule.professor.first_name} {class_schedule.professor.last_name}" if class_schedule.professor else "TBA",
            'course': class_schedule.course_section.course_name if class_schedule.course_section else "",
            'room': class_schedule.room_assignment or "TBA",
            'year_section': class_schedule.course_section.course_section if class_schedule.course_section else "",
            'schedule': f"{class_schedule.days} {class_schedule.time_in.strftime('%H:%M')}-{class_schedule.time_out.strftime('%H:%M')}",
        }
        
        # Get date ranges
        attendance_dates = AttendanceRecord.objects.filter(
            class_schedule=class_schedule
        ).values_list('date', flat=True).distinct().order_by('date')
        
        date_ranges = []
        dates_list = list(attendance_dates)
        for i in range(0, len(dates_list), 8):
            chunk = dates_list[i:i+8]
            if chunk:
                start_date = chunk[0]
                end_date = chunk[-1]
                date_ranges.append({
                    'value': f"{start_date.strftime('%Y-%m-%d')}_to_{end_date.strftime('%Y-%m-%d')}",
                    'label': f"{start_date.strftime('%b %d, %Y')} - {end_date.strftime('%b %d, %Y')}"
                })
        
        # Get attendance data if date range is selected
        attendance_table = []
        date_headers = []
        
        if date_range:
            try:
                start_str, end_str = date_range.split('_to_')
                start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
                end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
                
                dates_in_range = AttendanceRecord.objects.filter(
                    class_schedule=class_schedule,
                    date__range=[start_date, end_date]
                ).values_list('date', flat=True).distinct().order_by('date')
                
                date_headers = [d.strftime('%m/%d') for d in list(dates_in_range)[:8]]
                
                students = Account.objects.filter(
                    course_section=class_schedule.course_section,
                    role='Student'
                ).order_by('last_name', 'first_name')
                
                for student in students:
                    student_data = {
                        'name': f"{student.last_name}, {student.first_name}",
                        'sex': student.sex or '',
                        'dates': []
                    }
                    
                    for date in list(dates_in_range)[:8]:
                        try:
                            record = AttendanceRecord.objects.get(
                                class_schedule=class_schedule,
                                student=student,
                                date=date
                            )
                            student_data['dates'].append({
                                'time_in': record.time_in.strftime('%H:%M') if record.time_in else '',
                                'time_out': record.time_out.strftime('%H:%M') if record.time_out else '',
                                'status': record.status
                            })
                        except AttendanceRecord.DoesNotExist:
                            student_data['dates'].append({'time_in': '', 'time_out': '', 'status': ''})
                    
                    attendance_table.append(student_data)
                
            except (ValueError, TypeError):
                pass
        
        return JsonResponse({
            'class_data': class_data,
            'date_ranges': date_ranges,
            'date_headers': date_headers,
            'attendance_table': attendance_table,
            'student_count': len(attendance_table)
        })
        
    except ClassSchedule.DoesNotExist:
        return JsonResponse({'error': 'Class schedule not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
    
    # TEMPORARY! delete attendance record
@instructor_or_admin_required
@require_POST
def clear_attendance(request):
    """Clear all attendance records"""
    try:
        # Count records before deletion
        count = AttendanceRecord.objects.all().count()
        
        # Delete all attendance records
        AttendanceRecord.objects.all().delete()
        
        return JsonResponse({
            'status': 'success',
            'count': count,
            'message': f'Deleted {count} attendance records'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

# TEMPORARY! delete student accounts
@instructor_or_admin_required
@require_POST
def clear_students(request):
    """Delete all student accounts"""
    try:
        # Count students before deletion
        count = Account.objects.filter(role='Student').count()
        
        # Delete all student accounts
        Account.objects.filter(role='Student').delete()
        
        return JsonResponse({
            'status': 'success',
            'count': count,
            'message': f'Deleted {count} student accounts'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)